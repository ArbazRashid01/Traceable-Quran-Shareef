#!/usr/bin/env python3
"""
Traceable Quran — DOCX Exporter
Generates an editable Word document matching the 5-page HTML preview layout.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from generate import WORDS, MEANINGS, PAGES, ar

# ── COLOUR PALETTE ────────────────────────────────────────────────────────
GOLD        = RGBColor(0xC9, 0xA8, 0x4C)
DARK_BROWN  = RGBColor(0x1A, 0x0E, 0x04)
AMBER       = RGBColor(0xA0, 0x78, 0x30)
PARCHMENT   = RGBColor(0xF6, 0xEA, 0xD6)
BODY_TEXT   = RGBColor(0x2A, 0x1A, 0x06)
GRAY_TRACE  = RGBColor(0xCC, 0xC0, 0xA8)   # light gray for Arabic (traceable)

# ── HELPERS ───────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    """Set table cell background colour."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, sides=('top','bottom','left','right'),
                    color='C9A84C', sz='4'):
    """Add border to specific sides of a cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in sides:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    sz)
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def no_border(cell):
    """Remove all borders from a cell."""
    set_cell_border(cell,
                    sides=('top','bottom','left','right'),
                    color='FFFFFF', sz='0')

def rtl_para(para):
    """Mark a paragraph as RTL."""
    pPr = para._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)

def set_row_height(row, height_cm):
    tr   = row._tr
    trPr = tr.get_or_add_trPr()
    trH  = OxmlElement('w:trHeight')
    trH.set(qn('w:val'),  str(int(height_cm * 567)))  # EMU approx
    trH.set(qn('w:hRule'), 'atLeast')
    trPr.append(trH)

def set_col_width(cell, width_cm):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW  = OxmlElement('w:tcW')
    tcW.set(qn('w:w'),    str(int(width_cm * 567)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

def para_in_cell(cell, text, font_name, font_size_pt,
                 bold=False, italic=False,
                 color=None, align=WD_ALIGN_PARAGRAPH.CENTER,
                 rtl=False, space_before=0, space_after=0):
    para = cell.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after  = Pt(space_after)
    if rtl:
        rtl_para(para)
    run = para.add_run(text)
    run.font.name  = font_name
    run.font.size  = Pt(font_size_pt)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    if rtl:
        rPr = run._r.get_or_add_rPr()
        r   = OxmlElement('w:rtl')
        rPr.append(r)
    return para


# ── WORD CELL BUILDER ─────────────────────────────────────────────────────
def build_word_cell(cell, arabic, translit, meaning):
    """Fill a word cell with 3 layers: transliteration / Arabic / meaning."""
    cell.paragraphs[0].clear()          # remove default empty paragraph

    # 1. Transliteration — top, amber italic
    para_in_cell(cell, translit,
                 font_name='Georgia', font_size_pt=6,
                 italic=True, color=AMBER,
                 space_before=1, space_after=0)

    # 2. Thin separator line (simulated with a border paragraph)
    sep = cell.add_paragraph()
    sep.paragraph_format.space_before = Pt(0)
    sep.paragraph_format.space_after  = Pt(0)
    sep_fmt = sep._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '2')
    bottom.set(qn('w:color'), 'E4D0A8')
    pBdr.append(bottom)
    sep_fmt.append(pBdr)

    # 3. Arabic — gray (traceable), large, RTL
    para_in_cell(cell, arabic,
                 font_name='Times New Roman', font_size_pt=16,
                 color=GRAY_TRACE, rtl=True,
                 space_before=2, space_after=2)

    # 4. Thin separator
    sep2 = cell.add_paragraph()
    sep2.paragraph_format.space_before = Pt(0)
    sep2.paragraph_format.space_after  = Pt(0)
    s2f = sep2._p.get_or_add_pPr()
    pBdr2 = OxmlElement('w:pBdr')
    b2 = OxmlElement('w:bottom')
    b2.set(qn('w:val'),   'single')
    b2.set(qn('w:sz'),    '2')
    b2.set(qn('w:color'), 'E4D0A8')
    pBdr2.append(b2)
    s2f.append(pBdr2)

    # 5. English meaning — bottom, small
    para_in_cell(cell, meaning,
                 font_name='Georgia', font_size_pt=6,
                 color=BODY_TEXT,
                 space_before=0, space_after=1)

def build_marker_cell(cell, verse_key):
    """Fill an Ayah-end marker cell."""
    cell.paragraphs[0].clear()
    n = int(verse_key.split(':')[1])
    symbol = f'۝{ar(n)}'
    para_in_cell(cell, symbol,
                 font_name='Times New Roman', font_size_pt=12,
                 color=GOLD, rtl=True,
                 space_before=4, space_after=0)


# ── PAGE BUILDER ──────────────────────────────────────────────────────────
def build_page(doc, p):
    WPR      = p['wpr']
    is_first = (p['num'] == 1)

    # ── PAGE BREAK (not on page 1) ────────────────────────────────────────
    if not is_first:
        doc.add_page_break()

    # ── HEADER TABLE: JUZ | SURAH NAME | PAGE ─────────────────────────────
    hdr = doc.add_table(rows=1, cols=3)
    hdr.style = 'Table Grid'
    hdr.alignment = WD_TABLE_ALIGNMENT.CENTER
    juz_c, mid_c, pg_c = hdr.rows[0].cells

    # top/bottom gold borders on header row
    for c in (juz_c, mid_c, pg_c):
        set_cell_bg(c, 'F8F0DC')
        set_cell_border(c, ('top','bottom'), 'C9A84C', '12')
        set_cell_border(c, ('left','right'),  'FFFFFF', '0')

    set_col_width(juz_c, 3.0)
    set_col_width(mid_c, 9.0)
    set_col_width(pg_c,  3.0)

    para_in_cell(juz_c, f'JUZ {p["juz"]}',
                 'Georgia', 8, bold=True, color=RGBColor(0x7A,0x58,0x10),
                 align=WD_ALIGN_PARAGRAPH.LEFT)
    para_in_cell(mid_c, f'❖  {p["surah_en"]}  ❖',
                 'Georgia', 11, bold=True, color=DARK_BROWN)
    para_in_cell(pg_c, f'PAGE {p["num"]}',
                 'Georgia', 8, bold=True, color=RGBColor(0x7A,0x58,0x10),
                 align=WD_ALIGN_PARAGRAPH.RIGHT)

    # ── SURAH START BAND ──────────────────────────────────────────────────
    if p['new_surah']:
        sb = doc.add_table(rows=2, cols=1)
        sb.style = 'Table Grid'
        r1c = sb.rows[0].cells[0]
        r2c = sb.rows[1].cells[0]
        set_cell_bg(r1c, 'EDE0AA')
        set_cell_bg(r2c, 'F4E8C4')
        set_cell_border(r1c, ('top',),    'C9A84C', '8')
        set_cell_border(r2c, ('bottom',), 'C9A84C', '8')
        para_in_cell(r1c, p['surah_ar'],
                     'Times New Roman', 14, color=DARK_BROWN, rtl=True)
        para_in_cell(r2c,
                     f'Surah {p["surah_no"]}  •  {p["surah_meaning"]}'
                     f'  •  Revealed in {p["revealed"]}'
                     f'  •  {p["verses_total"]} Verses',
                     'Georgia', 8, color=RGBColor(0x7A,0x58,0x10))

    # ── BISMILLAH ─────────────────────────────────────────────────────────
    if p['bismillah_separate']:
        bm = doc.add_table(rows=2, cols=1)
        bm.style = 'Table Grid'
        bt_c = bm.rows[0].cells[0]
        ba_c = bm.rows[1].cells[0]
        for c in (bt_c, ba_c):
            set_cell_bg(c, 'FDF7ED')
            set_cell_border(c, ('top','bottom'), 'C9A84C', '4')
            set_cell_border(c, ('left','right'),  'FFFFFF', '0')
        para_in_cell(bt_c,
                     'Bis-mil-lā-hir  Raḥ-mā-nir  Raḥīm',
                     'Georgia', 7, italic=True, color=AMBER)
        para_in_cell(ba_c,
                     'بِسْمِ ٱللَّهِ ٱلرَّحْمَـٰنِ ٱلرَّحِيمِ',
                     'Times New Roman', 18, color=GRAY_TRACE, rtl=True)


    # ── BODY: 2-COL TABLE (meanings | word rows) ─────────────────────────
    body = doc.add_table(rows=1, cols=2)
    body.style = 'Table Grid'
    body.alignment = WD_TABLE_ALIGNMENT.CENTER
    left_col = body.rows[0].cells[0]
    right_col = body.rows[0].cells[1]

    set_col_width(left_col,  3.2)
    set_col_width(right_col, 11.8)
    set_cell_bg(left_col,  'F6EAD6')
    set_cell_bg(right_col, 'FDF7ED')
    set_cell_border(left_col, ('right',), 'D4B870', '6')

    # LEFT — compact verse meanings
    left_col.paragraphs[0].clear()
    for vk in p['verses']:
        vnum = ar(int(vk.split(':')[1]))
        meaning_text = MEANINGS[vk]
        mp = left_col.add_paragraph()
        mp.paragraph_format.space_before = Pt(3)
        mp.paragraph_format.space_after  = Pt(3)
        # verse number badge
        num_run = mp.add_run(f'{vnum} ')
        num_run.font.name  = 'Times New Roman'
        num_run.font.size  = Pt(9)
        num_run.font.bold  = True
        num_run.font.color.rgb = GOLD
        # meaning text
        txt_run = mp.add_run(f'"{meaning_text}"')
        txt_run.font.name   = 'Georgia'
        txt_run.font.size   = Pt(7.5)
        txt_run.font.italic = True
        txt_run.font.color.rgb = BODY_TEXT
        # dashed separator
        pPr = mp._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'),   'dashed')
        bot.set(qn('w:sz'),    '4')
        bot.set(qn('w:color'), 'D4B870')
        pBdr.append(bot)
        pPr.append(pBdr)

    # RIGHT — build all items in continuous flow
    items = []
    for vk in p['verses']:
        for i in range(len(WORDS[vk])):
            items.append(('W', vk, i))
        items.append(('M', vk, -1))

    wpr     = WPR
    chunks  = [items[i:i+wpr] for i in range(0, len(items), wpr)]
    nrows   = len(chunks)

    # nested table inside right_col
    right_col.paragraphs[0].clear()
    # We'll add one paragraph per row (simpler than nested tables)
    for ri, chunk in enumerate(chunks):
        # Row table: 1 row × N cells
        row_tbl = right_col.add_table(rows=1, cols=len(chunk))
        row_tbl.style = 'Table Grid'
        tr = row_tbl.rows[0]

        for ci, (t, vk, wi) in enumerate(reversed(chunk)):   # reversed = RTL
            cell = tr.cells[ci]
            no_border(cell)
            set_cell_border(cell, ('bottom',), 'E8D8B8', '2')
            if t == 'M':
                set_col_width(cell, 0.7)
                build_marker_cell(cell, vk)
            else:
                ar_text, tr_text, en_text = WORDS[vk][wi]
                build_word_cell(cell, ar_text, tr_text, en_text)

        # add small spacing between rows
        sp = right_col.add_paragraph()
        sp.paragraph_format.space_before = Pt(1)
        sp.paragraph_format.space_after  = Pt(0)

    # ── FOOTER ────────────────────────────────────────────────────────────
    ft = doc.add_table(rows=1, cols=2)
    ft.style = 'Table Grid'
    fl = ft.rows[0].cells[0]
    fr = ft.rows[0].cells[1]
    for c in (fl, fr):
        set_cell_border(c, ('top',),             'C9A84C', '12')
        set_cell_border(c, ('bottom','left','right'), 'FFFFFF', '0')
    para_in_cell(fl,
                 f'Surah {p["surah_en"]} — {p["surah_meaning"]}',
                 'Georgia', 8, color=RGBColor(0x7A,0x58,0x10),
                 align=WD_ALIGN_PARAGRAPH.LEFT)
    para_in_cell(fr,
                 f'Page {p["num"]}  |  Juz {p["juz"]}',
                 'Georgia', 8, color=RGBColor(0x7A,0x58,0x10),
                 align=WD_ALIGN_PARAGRAPH.RIGHT)


# ── DOCUMENT SETUP ────────────────────────────────────────────────────────
def build_docx():
    doc = Document()

    # A4 page size + narrow margins
    section = doc.sections[0]
    section.page_width   = Cm(21.0)
    section.page_height  = Cm(29.7)
    section.top_margin   = Cm(1.2)
    section.bottom_margin= Cm(1.0)
    section.left_margin  = Cm(1.2)
    section.right_margin = Cm(1.2)

    # Default paragraph spacing
    style = doc.styles['Normal']
    style.font.name = 'Georgia'
    style.font.size = Pt(9)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after  = Pt(0)

    for p in PAGES:
        print(f"  Building page {p['num']}: {p['surah_en']} "
              f"({len(p['verses'])} verses)...")
        build_page(doc, p)

    out = 'Traceable-Quran-Preview.docx'
    doc.save(out)
    print(f"\nSaved: {out}")
    return out

if __name__ == '__main__':
    print("Generating Traceable Quran DOCX...")
    build_docx()
